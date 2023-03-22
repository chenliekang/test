from win32com import client as wc
from win32com.client import gencache, Dispatch
import docx
import traceback
import re


def doSaveAas(doc_path):
    # doc文档另存为docx文档
    # 打开word应用程序
    # gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 1, bForDemand=True)
    word = wc.Dispatch('Word.Application')
    # wordapp = wc.gencache.EnsureDispatch("Word.Application")
    docx_path = doc_path + 'x'
    doc = word.Documents.Open(doc_path)        # 目标路径下的文件
    doc.SaveAs(docx_path, 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
    doc.Close()
    word.Quit()


def direct_deal_doc(doc_path):
    # gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 1, bForDemand=True)
    word = wc.Dispatch('Word.Application')
    word.Visible = 0
    word.DisplayAlerts = 0
    doc = word.Documents.Open(FileName=doc_path, Encoding='gbk')
    for para in doc.paragraphs:
        print(para.Range.Text)
    doc.Close()
    word.Quit


def get_value(content, start_tag, end_tag, other_tag=[], ignore_cl=False):
    """
    从pdf获取数据，返回的数据为list类型
    :param content: 字符串类型，获取数据的原文
    :param start_tag: 字符串类型，所需获取数据左边紧靠内容
    :param end_tag: 字符串类型，所需获取数据右边紧靠内容
    :param other_tag: list类型，可选，默认为空，所需获取数据的前提条件，越靠近数据的前提条件在list里越靠前
    :param ignore_cl: boolean类型，可选，默认为False，是否忽略所需获取数据的空格或换行符
    :return:
    """
    # 根据是否有前提条件确定正则表达式的字符串
    if other_tag:
        pattern_content = start_tag + '(.*?)' + end_tag
        for premise in other_tag:
            pattern_content = premise + '(.*?)' + pattern_content
    else:
        pattern_content = start_tag + '(.*?)' + end_tag
    # 根据是否忽略所需数据中的空格和换行符等确定正则表达式
    if ignore_cl:
        pattern = re.compile(r'{}'.format(pattern_content))
    else:
        pattern = re.compile(r'{}'.format(pattern_content), re.DOTALL)
    # 根据正则表达式获取数据
    values = pattern.findall(content)
    # 优化获取下来的数据（把数据中的空格去掉）
    new_values = []
    if other_tag:
        for value in values:
            new_values.append(re.sub(re.compile(r'\s+'), '', value[-1]))
    else:
        for value in values:
            new_values.append(re.sub(re.compile(r'\s+'), '', value))
    return new_values


def get_value2(content, start_tag, end_tag, value, other_tag=[]):
    if not value:
        txt = get_value(content, start_tag, end_tag, other_tag)
        return txt
    else:
        return value


def get_rent_info(text, info=None):
    info = {
        "合同甲方": [], "合同约定收款人": [], "合同约定收款账号": [], "合同的不含税金额": [], "合同支付方式": []
    }

    info["合同甲方"] += get_value(text, '甲方： ', ' ')   # 1 9 11 12 13 15 17 18
    info["合同甲方"] += get_value(text, '甲方：', ' ')  # 1 2 6 9 11 12 13 14 15 16 17 18
    info["合同甲方"] += get_value(text, '甲方（出租人）：', ' ')  # 7

    sign = ['签字：', '签字： ', '[']
    for i in sign:
        if i in info["合同甲方"]:
            info["合同甲方"].remove(i)
    while '' in info["合同甲方"]:
        info["合同甲方"].remove('')
    info["合同甲方"] = list(set(info["合同甲方"]))

    info["合同约定收款人"] = get_value2(text, '户名：\[', '\]', info["合同约定收款人"], ['账户信息'])  # 6 7 11 17 18
    info["合同约定收款人"] = get_value2(text, '甲方：\s*', ' ', info["合同约定收款人"], ['账户信息'])  # 1 2 9 13 14 15 16
    info["合同约定收款人"] = get_value2(text, '户名：\s*', ' ', info["合同约定收款人"], [''])  # 12
    while '' in info["合同约定收款人"]:
        info["合同约定收款人"].remove('')
    info["合同约定收款人"] = list(set(info["合同约定收款人"]))

    info["合同约定收款账号"] = get_value2(text, '账号：\[', '\]', info["合同约定收款账号"], ['账户信息'])  # 6 7 11 17 18
    info["合同约定收款账号"] = get_value2(text, '账\s*号', ' ', info["合同约定收款账号"], ['账户信息'])  # 1
    info["合同约定收款账号"] = get_value2(text, '帐\s*号', ' ', info["合同约定收款账号"], ['账户信息'])  # 2 9 13 14 15 16
    info["合同约定收款账号"] = get_value2(text, ' ', '，', info["合同约定收款账号"], ['卡号：'])  # 12
    while '' in info["合同约定收款账号"]:
        info["合同约定收款账号"].remove('')
    info["合同约定收款账号"] = list(set(info["合同约定收款账号"]))
    for i in range(len(info["合同约定收款账号"])):
        if '：' in info["合同约定收款账号"][i]:
            info["合同约定收款账号"][i] = info["合同约定收款账号"][i].strip('：')
    if len(info["合同约定收款账号"]) == 2:
        id1 = text.find(info["合同约定收款账号"][0])
        id2 = text.find(info["合同约定收款账号"][1])
        if id1 < id2:
            info["合同约定收款账号"] = info["合同约定收款账号"][0]
        else:
            info["合同约定收款账号"] = info["合同约定收款账号"][1]
        info["合同约定收款账号"] = [info["合同约定收款账号"]]

    # info["合同的不含税金额"] = get_value2(text, '小写 \[', ']', info["合同的不含税金额"], ['租金及支付方式'])#
    info["合同的不含税金额"] = get_value2(text, '不含税价格为', '元', info["合同的不含税金额"], ['租赁期限'])  # 6 7
    info["合同的不含税金额"] = get_value2(text, '不含税金额为：\[', '\]', info["合同的不含税金额"],
                                  ['租金的支付方式及相关费用的支付'])  # 1 2 9 13 14 15 16
    info["合同的不含税金额"] = get_value2(text, '小写\[', '\]', info["合同的不含税金额"], ['租赁期、租金及支付'])  # 11
    info["合同的不含税金额"] = get_value2(text, '小写\[', '\]', info["合同的不含税金额"], ['价款'])  # 17 18
    info["合同的不含税金额"] = get_value2(text, '租金', '元', info["合同的不含税金额"], ['第一次'])  # 12
    while '' in info["合同的不含税金额"]:
        info["合同的不含税金额"].remove('')
    info["合同的不含税金额"] = list(set(info["合同的不含税金额"]))

    info["合同支付方式"] = get_value2(text, '租金按\[', '\]', info["合同支付方式"], ['租赁期'])  # 6 7 11
    info["合同支付方式"] = get_value2(text, '每', '支付', info["合同支付方式"], ['交纳期限为'])  # 1 2 9 13 14 15 16
    # info["合同支付方式"] = get_value2(text, '支付方式为：每\[', '\]', info["合同支付方式"], ['租金及支付方式 '])#
    info["合同支付方式"] = get_value2(text, '支付方式：每\[', '\]', info["合同支付方式"], ['租金'])  # 17
    info["合同支付方式"] = get_value2(text, '租金按', '支付', info["合同支付方式"], ['租赁期'])  # 18
    while '' in info["合同支付方式"]:
        info["合同支付方式"].remove('')
    info["合同支付方式"] = list(set(info["合同支付方式"]))
    if len(info["合同支付方式"]) == 1:
        if '季度HESJA' in info["合同支付方式"][0]:
            info["合同支付方式"] = ['季度']

    return info


if __name__ == '__main__':
    doc_path = r'C:\Users\95111\Desktop\合同\石家庄电信（股份）2020年与石家庄万川投资有限公司市区南二环金利街东南角美化塔及附属设施租赁合同.doc'
    docx_path = doc_path + 'x'
    try:
        direct_deal_doc(doc_path)
        # doSaveAas(doc_path)
        file = docx.Document(docx_path)
        # print("段落数:" + str(len(file.paragraphs)))
        text = ''
        for para in file.paragraphs:
            text += para.text
            text += ' '
            print(para.text)

        info = {}
        info = get_rent_info(text)
        pass
    except Exception as e:
        traceback.print_exc()
        pass